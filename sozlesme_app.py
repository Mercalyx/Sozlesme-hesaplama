import streamlit as st
import pandas as pd
import io
from docx import Document  # <-- yeni eklenecek

def replace_room_table(doc: Document, konaklama_data: list):  # <-- tamamı buraya yapıştır
    for i, paragraph in enumerate(doc.paragraphs):
        if "{{room_table}}" in paragraph.text:
            paragraph.text = ""
            table = doc.add_table(rows=1, cols=5)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Tarih"
            hdr_cells[1].text = "Oda Türü"
            hdr_cells[2].text = "Oda Sayısı"
            hdr_cells[3].text = "Gecelik Fiyat"
            hdr_cells[4].text = "Toplam Tutar"

            for row in konaklama_data:
                tarih = str(row.get("Tarih", ""))
                oda_turu = str(row.get("Oda Türü", ""))
                sayi = int(row.get("Oda Sayısı", 0))
                fiyat = float(row.get("Gecelik Fiyat", 0))
                toplam = sayi * fiyat

                row_cells = table.add_row().cells
                row_cells[0].text = tarih
                row_cells[1].text = oda_turu
                row_cells[2].text = str(sayi)
                row_cells[3].text = f"{fiyat:.2f}"
                row_cells[4].text = f"{toplam:.2f}"

                doc._body._body.insert(i + 1, table._tbl)
                break

st.markdown("---")

st.set_page_config(page_title="Sözleşme Hesaplama Robotu", page_icon="🧾")

st.title("🧾 Sözleşme Tutarı Hesaplama Robotu")
st.markdown("---")

# Şirket/Acenta bilgileri
st.header("👤 Şirket / Acenta Bilgileri")
misafir_adi = st.text_input("Şirket veya Acenta Adı", "")
Tursab = st.radio("TURSAB Üyesi mi?", ["Hayır", "Evet"])

st.markdown("---")

# Para Birimi Seçimi
st.header("💱 Para Birimi Seçimi")
para_birimleri = {"EUR": "€", "USD": "$", "TL": "₺"}
secili_para_birimi = st.selectbox("Para Birimi Seçin", list(para_birimleri.keys()))
sembol = para_birimleri[secili_para_birimi]

st.markdown("---")

# Veri Giriş Yöntemi Seçimi
st.header("📋 Veri Giriş Yöntemi Seçimi")
giris_yontemi = st.radio(
    "Verileri nasıl gireceksiniz?",
    ("Dosya Yükleyerek", "Tabloyla Giriş",  "Manuel Giriş")
)

st.markdown("---")

oda_bilgileri = []
etkinlikler = []

# Eğer Dosya Yüklenirse
if giris_yontemi == "Dosya Yükleyerek":
    st.header("📂 Booking Verisi Yükle")
    uploaded_file = st.file_uploader("Booking Verisini Yükleyin (CSV formatında)", type=["csv"])

    if uploaded_file is not None:
        df = pd.read_csv(uploaded_file)
        st.success("✅ Dosya Başarıyla Yüklendi!")

        max_gun = df['Gün'].max()

        for gun in range(1, max_gun+1):
            gunluk_oda = df[df['Gün'] == gun].iloc[0]
            oda_bilgileri.append({
                "tek": gunluk_oda['Tek Kişilik Oda Sayısı'],
                "cift": gunluk_oda['Çift Kişilik Oda Sayısı'],
                "tek_f": gunluk_oda['Tek Kişilik Fiyat'],
                "cift_f": gunluk_oda['Çift Kişilik Fiyat'],
            })

            gunluk_etkinlikler = []
            gun_df = df[df['Gün'] == gun]
            for idx, row in gun_df.iterrows():
                gunluk_etkinlikler.append({
                    "tur": row['Etkinlik Türü'],
                    "kisi": row['Katılımcı Sayısı'],
                    "fiyat": row['Etkinlik Fiyatı']
                })
            etkinlikler.append(gunluk_etkinlikler)

elif giris_yontemi == "Tabloyla Giriş":
    st.header("📅 Etkinlik Tablosu")
    etkinlik_df = pd.DataFrame({
        "Tarih": [""],
        "Etkinlik Türü": [""],
        "Katılımcı Sayısı": [0],
        "Kişi Başı Fiyat": [0.0]
    })
    etkinlik_input = st.data_editor(etkinlik_df, num_rows="dynamic", use_container_width=True)

    st.header("🛏 Konaklama Tablosu")
    konaklama_df = pd.DataFrame({
        "Tarih": [""],
        "Oda Türü": [""],  # "Tek" ya da "Çift"
        "Oda Sayısı": [0],
        "Gecelik Fiyat": [0.0]
    })
    konaklama_input = st.data_editor(konaklama_df, num_rows="dynamic", use_container_width=True)

    # Etkinlik verisini işle
    etkinlikler = []
    grouped = etkinlik_input.groupby("Tarih")
    for tarih, grup in grouped:
        gunluk_etkinlikler = []
        for _, row in grup.iterrows():
            gunluk_etkinlikler.append({
                "tur": row["Etkinlik Türü"],
                "kisi": int(row["Katılımcı Sayısı"]),
                "fiyat": float(row["Kişi Başı Fiyat"])
            })
        etkinlikler.append(gunluk_etkinlikler)

    # Oda verisini işle
    oda_bilgileri = []
    grouped_konaklama = konaklama_input.groupby("Tarih")
    for tarih, grup in grouped_konaklama:
        tek = 0
        cift = 0
        tek_f = 0.0
        cift_f = 0.0
        for _, row in grup.iterrows():
            if row["Oda Türü"] == "Tek":
                tek += int(row["Oda Sayısı"])
                tek_f = float(row["Gecelik Fiyat"])
            elif row["Oda Türü"] == "Çift":
                cift += int(row["Oda Sayısı"])
                cift_f = float(row["Gecelik Fiyat"])
        oda_bilgileri.append({
            "tek": tek,
            "cift": cift,
            "tek_f": tek_f,
            "cift_f": cift_f
})

# Eğer Manuel Giriş Seçilirse
else:
    # Süre Bilgileri
    st.header("🗓 Süre Bilgileri")
    etkinlik_gun_sayisi = st.number_input("Etkinlik Süresi (Gün)", min_value=1, value=1)
    konaklama_gun_sayisi = st.number_input("Konaklama Süresi (Gece)", min_value=1, value=1)

    st.markdown("---")

    # Farklılık seçenekleri
    st.header("⚙ Bilgi Giriş Ayarları")
    etkinlik_farkli_mi = st.checkbox("Her gün için farklı etkinlik bilgisi girilsin mi?", value=True)
    oda_farkli_mi = st.checkbox("Her gün için farklı oda bilgisi girilsin mi?", value=True)

    etkinlik_fiyat_degisim = st.radio("Her gün etkinlik fiyatı değişiyor mu?", ["Hayır", "Evet"]) if etkinlik_farkli_mi else "Hayır"
    oda_fiyat_degisim = st.radio("Her gün oda fiyatı değişiyor mu?", ["Hayır", "Evet"]) if oda_farkli_mi else "Hayır"

    st.markdown("---")

    # Başlangıç Fiyatları
    st.header("💶 Başlangıç Oda ve Etkinlik Fiyatları")
    tek_kisilik_standart_fiyat = st.number_input("Tek Kişilik Oda Standart Fiyatı (gecelik)", min_value=0.0, value=0.0)
    cift_kisilik_standart_fiyat = st.number_input("Çift Kişilik Oda Standart Fiyatı (gecelik)", min_value=0.0, value=0.0)

    etkinlik_turleri = ["Toplantı", "Gala", "Kokteyl", "Öğle Yemeği", "Akşam Yemeği", "Breakout", "Kurulum"]
    standart_etkinlik_fiyatlari = {}
    for tur in etkinlik_turleri:
        fiyat = st.number_input(f"{tur} Standart Fiyatı (Kişi Başı)", min_value=0.0, value=0.0, key=f"standart_{tur}")
        standart_etkinlik_fiyatlari[tur] = fiyat

    st.markdown("---")

    # Oda Bilgileri
    st.header("🛏 Konaklama Bilgileri")
    for gun in range(konaklama_gun_sayisi):
        st.subheader(f"{gun+1}. Gece Oda Bilgisi")
        tek = st.number_input(f"Tek Kişilik Oda Sayısı (Gece {gun+1})", min_value=0, key=f"tek{gun}")
        cift = st.number_input(f"Çift Kişilik Oda Sayısı (Gece {gun+1})", min_value=0, key=f"cift{gun}")

        if oda_farkli_mi and oda_fiyat_degisim == "Evet":
            tek_f = st.number_input(f"Tek Kişilik Oda Fiyatı (Gece {gun+1})", min_value=0.0, key=f"tekf{gun}")
            cift_f = st.number_input(f"Çift Kişilik Oda Fiyatı (Gece {gun+1})", min_value=0.0, key=f"ciftf{gun}")
        else:
            tek_f = tek_kisilik_standart_fiyat
            cift_f = cift_kisilik_standart_fiyat

        oda_bilgileri.append({"tek": tek, "cift": cift, "tek_f": tek_f, "cift_f": cift_f})

    st.markdown("---")

    # Etkinlik Bilgileri
    st.header("🎤 Etkinlik Bilgileri")
    if "etkinlik_sayaci" not in st.session_state:
        st.session_state.etkinlik_sayaci = {}

    for gun in range(etkinlik_gun_sayisi):
        st.subheader(f"{gun+1}. Gün Etkinlikleri")

        if gun not in st.session_state.etkinlik_sayaci:
            st.session_state.etkinlik_sayaci[gun] = 1

        if st.button(f"{gun+1}. Gün İçin Etkinlik Ekle", key=f"etkinlik_ekle_btn_{gun}"):
            st.session_state.etkinlik_sayaci[gun] += 1

        g_etkinlikler = []
        for j in range(st.session_state.etkinlik_sayaci[gun]):
            tur = st.selectbox(f"Etkinlik Türü {j+1} (Gün {gun+1})", options=etkinlik_turleri, key=f"tur_{gun}_{j}")

            if etkinlik_farkli_mi and etkinlik_fiyat_degisim == "Evet":
                fiyat = st.number_input(f"{tur} Fiyatı (Gün {gun+1})", min_value=0.0, key=f"fiyat_{gun}_{j}")
            else:
                fiyat = standart_etkinlik_fiyatlari.get(tur, 0)

            kisi = st.number_input(f"{tur} Katılımcı Sayısı (Gün {gun+1})", min_value=0, key=f"kisi_{gun}_{j}")
            g_etkinlikler.append({"tur": tur, "fiyat": fiyat, "kisi": kisi})

        etkinlikler.append(g_etkinlikler)

st.markdown("---")

# Hesaplama Modülü
st.header("🧮 Hesaplama")

vergisiz_konaklama = sum([gun_bilgi["tek"] * gun_bilgi["tek_f"] + gun_bilgi["cift"] * gun_bilgi["cift_f"] for gun_bilgi in oda_bilgileri])
vergili_konaklama = vergisiz_konaklama * 1.12
vergi_konaklama = vergisiz_konaklama * 0.12

vergisiz_etkinlik = 0
vergi_etkinlik = 0
vergili_etkinlik = 0
for gun in etkinlikler:
    for etkinlik in gun:
        e_tutar = etkinlik["fiyat"] * etkinlik["kisi"]
        vergi = e_tutar * 0.20
        vergisiz_etkinlik += e_tutar
        vergi_etkinlik += vergi
        vergili_etkinlik += e_tutar + vergi

damga_vergisi = 0
if Tursab == "Evet":
    damga_vergisi = 3783.20 / 82 + vergisiz_etkinlik * 0.00948 / 2
else:
    damga_vergisi = (vergisiz_konaklama + vergisiz_etkinlik) * 0.00948 / 2

toplam_tutar = vergili_konaklama + vergili_etkinlik + damga_vergisi
ilk_odeme = toplam_tutar * 0.30
son_odeme = toplam_tutar * 0.70

# Çıktı
st.success("✅ Hesaplama Tamamlandı!")

st.header("📋 Sözleşme Özeti")
st.info(f"Şirket/Acenta: {misafir_adi}")
st.write(f"Konaklama Bedeli (KDV Hariç): {vergisiz_konaklama:,.2f} {sembol}")
st.write(f"Etkinlik Bedeli (KDV Hariç): {vergisiz_etkinlik:,.2f} {sembol}")
st.write(f"Konaklama Bedeli (KDV Dahil): {vergili_konaklama:,.2f} {sembol}")
st.write(f"Etkinlik Bedeli (KDV Dahil): {vergili_etkinlik:,.2f} {sembol}")
st.write(f"Toplam KDV: {(vergi_konaklama + vergi_etkinlik):,.2f} {sembol}")
st.write(f"Damga Vergisi: {damga_vergisi:,.2f} {sembol}")
st.write(f"Genel Toplam: {toplam_tutar:,.2f} {sembol}")
st.subheader(f"🔵 İlk Ödeme (30%): {ilk_odeme:,.2f} {sembol}")
st.subheader(f"🔵 Kalan Ödeme (70%): {son_odeme:,.2f} {sembol}")

st.markdown("---")

if giris_yontemi == "Tabloyla Giriş":
    st.markdown("### 📄 Word Sözleşmesi Oluştur")

    if st.button("Sözleşmeyi Word Formatında Oluştur"):
        try:
            doc = Document("1- Standard Agreement - TR.docx")
            replace_room_table(doc, konaklama_input.to_dict(orient="records"))
            doc.save("sozlesme_dolu.docx")

            with open("sozlesme_dolu.docx", "rb") as f:
                st.download_button(
                    label="📥 Word Olarak İndir",
                    data=f,
                    file_name="sozlesme_dolu.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
        except Exception as e:
            st.error(f"Sözleşme oluşturulurken hata oluştu: {e}")

st.markdown("---")

# Excel Export
st.header("📄 Teklif Excel Çıktısı")
data = {
    "Şirket/Acenta": [misafir_adi],
    "Konaklama (KDV Hariç)": [vergisiz_konaklama],
    "Etkinlik (KDV Hariç)": [vergisiz_etkinlik],
    "Konaklama (KDV Dahil)": [vergisiz_konaklama * 1.12],
    "Etkinlik (KDV Dahil)": [vergisiz_etkinlik * 1.20],
    "Toplam KDV": [vergi_konaklama + vergi_etkinlik],
    "Damga Vergisi": [damga_vergisi],
    "Genel Toplam": [toplam_tutar],
    "İlk Ödeme (30%)": [ilk_odeme],
    "Kalan Ödeme (70%)": [son_odeme]
}

df = pd.DataFrame(data)

@st.cache_data
def convert_df_to_excel(df):
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='Sözleşme Özeti')
    processed_data = output.getvalue()
    return processed_data

excel_data = convert_df_to_excel(df)

st.download_button(
    label="📥 Excel Olarak İndir",
    data=excel_data,
    file_name='sozlesme_ozeti.xlsx',
    mime='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
)

st.success("✅ Teklif Excel dosyasını indirebilirsiniz!")
